import sys
import sqlite3
import datetime as dt
from docx import Document
from calendar import monthrange
from collections import defaultdict
import webbrowser

from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import QApplication, QHeaderView, QTableWidgetItem
from PyQt5.QtWidgets import QMainWindow, QListWidgetItem, QFileDialog, QMessageBox

from BaseWindowClass import BaseWindow

from NotesMenu import Ui_NotesMenu
from NoteWindow import Ui_NoteWindow
from ToDoMenu import Ui_ToDoMenu
from ToDoWindow import Ui_ToDoWindow
from SettingsWindow import Ui_SettingsWindow
from CalendarWindow import Ui_CalendarWindow
from InfoWindow import Ui_InfoWindow


def switch_to_notes():
    notes.show()
    notes.setGeometry(todos.geometry())
    todos.hide()


def switch_to_todo():
    todos.show()
    todos.setGeometry(notes.geometry())
    notes.hide()


def update_notes_list():
    notes.update_list()


def set_settings(refresh=False):
    if refresh:
        settings.refresh()
    todos.set_settings(*settings.params)


def get_new_todo_id():
    return todos.get_new_todo_id()


class NotesMenu(BaseWindow, Ui_NotesMenu):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.con = sqlite3.connect("notes_db.sqlite")

        # свернуть и закрыть - функции кнопок верхней панели
        self.closeButton.clicked.connect(self.close)
        # self.closeButton.mouseMoveEvent = self.closeButton.setStyleSheet("background-color: red")
        self.hideButton.clicked.connect(self.showMinimized)
        self.titleBar.mousePressEvent = self.titleBarMousePressEvent

        # события для кнопок слева
        self.goToToDo.mousePressEvent = self.switch_to_todo
        self.addButton.mousePressEvent = self.add_note
        self.removeButton.mousePressEvent = self.remove_note
        self.exportButton.mousePressEvent = self.export_note
        self.listWidget.itemDoubleClicked.connect(self.note_clicked)
        self.askButton.mousePressEvent = self.open_info

        self.elements_dictionary = {}
        self.update_list()
        # ключ - номер в строчки в QListWidget, значение - id стикера

    def switch_to_todo(self, event):
        switch_to_todo()

    def add_note(self, event):  # добавляем стикер
        # подготавливаем id для нового стикера
        if self.elements_dictionary:
            max_id = max(self.elements_dictionary.values())

        else:
            max_id = 0

        # открываем окно для создания стикера и помечаем её как новую
        note = NoteWindow(max_id + 1, True)
        note.show()

    def remove_note(self, event):
        # удаляем стикер

        list_items = [x.row() for x in self.listWidget.selectedIndexes()]

        if not list_items:
            return None

        qm = QMessageBox()

        qm.setWindowTitle('Подтверждение удаления')
        qm.setText('Вы действительно хотите удалить выделенный стикер?')
        qm.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        qm.setStyleSheet("background-color: rgb(23, 28, 37); color: rgb(255, 255, 255)")
        button_yes = qm.button(QMessageBox.Yes)
        button_yes.setText('Да')
        button_no = qm.button(QMessageBox.No)
        button_no.setText('Нет')
        qm.exec_()

        if qm.clickedButton() == button_yes:  # необходимо подтвердить удаление
            note_id = self.elements_dictionary[list_items[0]]
            cur = self.con.cursor()
            cur.execute("DELETE FROM notes WHERE id = ?", (note_id, ))
            self.con.commit()
            self.update_list()

        else:
            qm.close()

    def export_note(self, event):  # экспортируем выбраный стикер в файл .txt
        if self.listWidget.selectedIndexes():
            file_name = QFileDialog.getSaveFileName(
                self, 'Сохранить в текстовый документ', '/',
                'Текстовый документ (*.txt)')[0]
            if file_name:
                cur = self.con.cursor()
                file = open(file_name, 'w')
                text = cur.execute(
                    "SELECT text FROM notes WHERE id = ?",
                    (self.elements_dictionary[self.listWidget.selectedIndexes()[0].row()],)
                ).fetchone()[0]
                file.write(text)
                file.close()

    def update_list(self):
        self.listWidget.clear()  # очищаем список стикеров, чтобы заполинть его
        cur = self.con.cursor()
        # собираем данные из обновленной БД
        ids = cur.execute("SELECT id FROM notes").fetchall()
        names = cur.execute("SELECT name FROM notes").fetchall()

        for i, (note_id, note_name) in enumerate(zip(ids, names)):
            self.elements_dictionary[i] = note_id[0]  # актуализируем словарь

            #  обновляем список стикеров
            self.listWidget.addItem(QListWidgetItem(str(note_name[0])))

    def note_clicked(self):  # пользователь решил отредактировать стикер
        # находим id стикера по её порядковому номеру
        note_id = self.elements_dictionary[self.listWidget.currentRow()]

        # открываем окно для редактирования стикера и помечаем её как не новую
        note = NoteWindow(note_id, False)
        note.show()

    def open_info(self, event):
        info = InfoWindow()
        info.show()


class NoteWindow(QMainWindow, Ui_NoteWindow):
    def __init__(self, note_id, is_new_note):
        super().__init__()
        self.setupUi(self)
        self.id = note_id

        self.setWindowFlags(Qt.Window | Qt.CustomizeWindowHint)
        self.drag_active = False  # флаг для drag and drop стикера
        self.spacer.mousePressEvent = self.spacer_press_event

        self.closeButton.mousePressEvent = self.edit_note

        self.con = sqlite3.connect("notes_db.sqlite")
        cur = self.con.cursor()

        if not is_new_note:  # редактируем уже имеющийся стикер

            self.nameInput.setText(cur.execute(
                "SELECT name FROM notes WHERE id = ?", (note_id,)).fetchone()[0])

            self.textEdit.setText(cur.execute(
                "SELECT text FROM notes WHERE id=?", (note_id,)).fetchone()[0])

        else:  # создаём новый стикер
            self.nameInput.setText("Новый стикер")

            cur.execute(
                "INSERT INTO notes(id, name, text) VALUES(?, ?, ?)",
                (self.id, self.nameInput.text(), self.textEdit.toPlainText()))
            self.con.commit()

            update_notes_list()

    def spacer_press_event(self, e):  # метод для drag and drop: запоминаем изначальную позицию
        self.previous_pos = e.globalPos()
        self.drag_active = True

    def mouseMoveEvent(self, e):  # метод для drag and drop: меняем позицию стикера
        if self.drag_active:
            delta = e.globalPos() - self.previous_pos
            self.move(self.x() + delta.x(), self.y() + delta.y())
            self.previous_pos = e.globalPos()

    def mouseReleaseEvent(self, e):
        if self.drag_active:
            self.drag_active = False

    def edit_note(self, event):
        cur = self.con.cursor()

        cur.execute('''UPDATE notes
        SET name = ?, text = ?
        WHERE id = ?''', (self.nameInput.text(), self.textEdit.toPlainText(), self.id))
        self.con.commit()

        update_notes_list()  # обновляем графическое представление БД
        self.close()


class ToDoMenu(BaseWindow, Ui_ToDoMenu):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        # self.titleBar, self.closeButton = self.titleBar, self.closeButton

        self.con = sqlite3.connect("notes_db.sqlite")

        # свернуть и закрыть - функции кнопок верхней панели
        self.closeButton.clicked.connect(self.close)
        # self.closeButton.mouseHoverEvent = self.closeButton.setStyleSheet("background-color: red")
        self.hideButton.clicked.connect(self.showMinimized)
        self.titleBar.mousePressEvent = self.titleBarMousePressEvent

        # события для кнопок слева
        self.goToNotes.mousePressEvent = self.switch_to_notes
        self.addButton.mousePressEvent = self.add_todo
        self.removeButton.mousePressEvent = self.remove_todo
        self.exportButton.mousePressEvent = self.export_todo
        self.tableWidget.itemClicked.connect(self.todo_clicked)
        self.tableWidget.itemDoubleClicked.connect(self.todo_double_clicked)
        self.calendarButton.mousePressEvent = self.open_calender
        self.settingsButton.mousePressEvent = self.open_settings
        self.askButton.mousePressEvent = self.open_info

        self.elements_dictionary = {}
        # ключ - номер в строчки в QTableWidget, значение - id стикера
        self.tableWidget.setColumnWidth(0, 5)
        header = self.tableWidget.horizontalHeader()
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.set_settings(*[None] * 4)

    def switch_to_notes(self, event):
        switch_to_notes()

    def add_todo(self, event):
        max_id = self.get_new_todo_id()

        # открываем окно для создания плана
        todo = ToDoWindow(max_id + 1, True)
        todo.show()

    def remove_todo(self, event):
        # удаляем план

        list_items = [x.row() for x in self.tableWidget.selectedIndexes()]

        if list_items:

            qm = QMessageBox()

            qm.setWindowTitle('Подтверждение удаления')
            qm.setText('Вы действительно хотите удалить выделенный план?')
            qm.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            qm.setStyleSheet("background-color: rgb(23, 28, 37); color: rgb(255, 255, 255)")
            button_yes = qm.button(QMessageBox.Yes)
            button_yes.setText('Да')
            button_no = qm.button(QMessageBox.No)
            button_no.setText('Нет')
            qm.exec_()

            if qm.clickedButton() == button_yes:  # необходимо подтвердить удаление
                todo_ids = [self.elements_dictionary[elem] for elem in list_items]
                cur = self.con.cursor()
                cur.execute("DELETE FROM todos WHERE id IN (" + ", ".join(
                    '?' * len(todo_ids)) + ")", todo_ids)
                self.con.commit()
                set_settings()

            else:
                qm.close()
        else:
            qm = QMessageBox()

            qm.setWindowTitle('Подтверждение удаления')
            qm.setText('Вы действительно хотите удалить все выполненные планы?')
            qm.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            qm.setStyleSheet("background-color: rgb(23, 28, 37); color: rgb(255, 255, 255)")
            button_yes = qm.button(QMessageBox.Yes)
            button_yes.setText('Да')
            button_no = qm.button(QMessageBox.No)
            button_no.setText('Нет')
            qm.exec_()

            if qm.clickedButton() == button_yes:  # необходимо подтвердить удаление
                cur = self.con.cursor()
                cur.execute("DELETE FROM todos WHERE done = 1")
                self.con.commit()
                set_settings()

            else:
                qm.close()

    def export_todo(self, event):  # экспорт текущего листа планов в .docx
        file_name = QFileDialog.getSaveFileName(
            self, 'Сохранить в текстовый документ', '/',
            'Текстовый документ (*.docx)')[0]
        document = Document()
        if file_name:
            cur = self.con.cursor()
            ids = tuple(self.elements_dictionary.values())
            #  ищем необходимую дедлайны и названия планов в базе данных
            information = list(cur.execute(
                "SELECT date, name FROM todos WHERE id IN"
                " (" + ", ".join('?' * len(ids)) + ")", ids).fetchall())

            without_dates = []
            dates = defaultdict(list)

            #  разделяем имеющиеся планы на категории по датам
            for info in information:
                date = info[0]
                if date:
                    dates[dt.date.fromisoformat(date)].append(info[1])
                else:
                    without_dates.append(info[1])  # отдельный список для планов без дедлайнов
            # формируем документ
            # он будет представлять из себя ненумерованный список, в котором
            # сначала идут планы без дедлайна, а потом сгруппированные по датам
            # и по возрастанию планы с дедлайнами
            for todo in without_dates:
                document.add_paragraph(todo, style='List Bullet')

            for date in sorted(dates.keys()):
                p = document.add_paragraph()
                p.add_run(str(date)).bold = True
                for name in dates[date]:
                    document.add_paragraph(name, style='List Bullet')

            document.save(file_name)

    def todo_clicked(self):  # меняем статус готовности плана
        if self.tableWidget.currentColumn() == 0:
            cur = self.con.cursor()
            current_id = self.elements_dictionary[self.tableWidget.currentRow()]
            current_done = cur.execute(
                "SELECT done FROM todos WHERE id = ?",
                (current_id, )).fetchone()[0]
            cur.execute(
                '''UPDATE todos
                SET done = ?
                WHERE id = ?''', (not current_done, current_id))
            self.con.commit()

            set_settings()

    def todo_double_clicked(self):  # открываем существующий план
        # находим id стикера по её порядковому номеру
        if self.tableWidget.currentColumn() == 1:
            todo_id = self.elements_dictionary[self.tableWidget.currentRow()]

            # открываем окно для редактирования стикера и помечаем её как не новую
            todo = ToDoWindow(todo_id, False)
            todo.show()

    def open_calender(self, event):  # открываем окно с календарём
        calendar = CalendarWindow()
        calendar.show()

    def open_settings(self, event):  # открываем окно с фильтром
        filter_window = SettingsWindow()
        filter_window.show()

    def set_settings(self, color_index, has_deadline, is_undone, date):
        # применяем настройки
        # просматриваем каждый из переданных параметров и формируем
        # из них запрос
        queries = []
        if color_index:
            queries.append(f'color_id = {color_index}')

        if has_deadline:
            queries.append('date IS NOT NULL')

        if is_undone:
            queries.append('done = 0')

        if date:
            queries.append(f'date = "{date.toString("yyyy-MM-dd")}"')

        self.tableWidget.clear()
        cur = self.con.cursor()

        if queries:  # если есть какие-то парамеры для фильтрации

            self.tableWidget.clear()
            cur = self.con.cursor()
            ids = cur.execute(
                f"SELECT id FROM todos WHERE {' AND '.join(queries)}").fetchall()
            selected_ids = tuple([todo_id[0] for todo_id in ids])
            todo_inf = cur.execute(
                "SELECT name, color_id, date, done FROM todos WHERE id IN (" +
                ", ".join('?' * len(selected_ids)) + ")", selected_ids).fetchall()

            colors_dict = {}
            if not color_index:
                # отсматриваем имеющиеся в базе данных цвета и устанавливаем
                # соответствия между id и кодами цветов
                colors = [information[1] for information in todo_inf]
                for color_id in colors:
                    color_code = cur.execute(
                        "SELECT color_code FROM colors WHERE color_id = ?", (color_id,)).fetchone()[
                        0]
                    colors_dict[color_id] = color_code
            else:
                color_code = cur.execute(
                    "SELECT color_code FROM colors WHERE color_id = ?",
                    (color_index,)).fetchone()[0]
                colors_dict[color_index] = color_code

        else:
            ids = cur.execute("SELECT id FROM todos").fetchall()
            todo_inf = cur.execute("SELECT name, color_id, date, done FROM todos").fetchall()
            colors = [information[1] for information in todo_inf]
            colors_dict = {}
            for color_id in colors:
                color_code = cur.execute(
                    "SELECT color_code FROM colors WHERE color_id = ?", (color_id,)).fetchone()[
                    0]
                colors_dict[color_id] = color_code

        if ids:  # если имеются какие-либо результаты запроса, выводим их
            self.tableWidget.setRowCount(len(ids))
            self.tableWidget.setColumnCount(2)
            for i, (todo_id, (todo_name, todo_color, todo_date, todo_done)) \
                    in enumerate(zip(ids, todo_inf)):

                self.elements_dictionary[i] = todo_id[0]
                background_color = list(map(int, colors_dict[todo_color].split(', ')))
                if todo_date is None:
                    todo_date = ''

                name, done = QTableWidgetItem(f"{todo_date} {todo_name}"), QTableWidgetItem()
                name.setBackground(QColor(*background_color))
                self.tableWidget.setItem(i, 1, name)

                if todo_done:
                    done.setText('a')
                else:
                    done.setText('r')

                done.setBackground(QColor(*background_color))
                done.setFont(QFont("Webdings"))
                self.tableWidget.setItem(i, 0, done)

    def get_new_todo_id(self):  # формируем уникальный id т.к. автоинкремент не работает
        ids = self.con.cursor().execute(
            "SELECT id FROM todos").fetchall()
        if self.elements_dictionary:
            return max(list([todo_id[0] for todo_id in ids])) + 1
        else:
            return 0

    def open_info(self, event):
        info = InfoWindow()
        info.show()


class ToDoWindow(BaseWindow, Ui_ToDoWindow):
    def __init__(self, todo_id, is_new_note, date=None):
        super().__init__()
        self.setupUi(self)
        self.id = todo_id
        self.is_new = is_new_note
        
        self.drag_active = False  # флаг для drag and drop плана
        self.spacer.mousePressEvent = self.titleBarMousePressEvent

        self.doneButton.mousePressEvent = self.edit_todo
        self.goBackButton.mousePressEvent = self.go_back
        self.colorBox.currentIndexChanged.connect(self.change_color)
        self.deadlineBox.stateChanged.connect(self.change_date)

        self.con = sqlite3.connect("notes_db.sqlite")
        cur = self.con.cursor()

        if not is_new_note:  # если план не новый, подгружаем в окно данные из БД

            self.nameInput.setText(cur.execute(
                "SELECT name FROM todos WHERE id = ?", (todo_id,)).fetchone()[0])

            self.infoEdit.setText(cur.execute(
                "SELECT info FROM todos WHERE id=?", (todo_id,)).fetchone()[0])

            color_id = cur.execute(
                "SELECT color_id FROM todos WHERE id = ?", (todo_id,)).fetchone()[0]
            self.colorBox.setCurrentIndex(color_id - 1)
            self.change_color()

            date = cur.execute(
                "SELECT date FROM todos WHERE id=?",  (todo_id,)).fetchone()[0]
            if date:
                self.deadlineBox.setChecked(True)
                self.dateEdit.setDate(dt.date.fromisoformat(date))
            else:
                self.dateEdit.setDate(dt.date.today())

        else:  # иначе создаём новый план
            self.nameInput.setText("Новый план")

            if date:
                self.deadlineBox.setChecked(True)
                self.dateEdit.setDate(date)
                cur.execute(
                    "INSERT INTO todos(id, name, info, date, color_id, done)"
                    " VALUES(?, ?, ?, ?, ?, ?)",
                    (self.id, self.nameInput.text(), self.infoEdit.toPlainText(),
                     date.toString("yyyy-MM-dd"), 1, False))
            else:
                self.dateEdit.setDate(dt.date.today())
                cur.execute(
                    "INSERT INTO todos(id, name, info, color_id, done) VALUES(?, ?, ?, ?, ?)",
                    (self.id, self.nameInput.text(), self.infoEdit.toPlainText(),
                     1, False))
            self.con.commit()
            set_settings()

    def edit_todo(self, event):  # сохраняем изменения в плане
        cur = self.con.cursor()

        if self.deadlineBox.isChecked():
            date = self.dateEdit.date().toString("yyyy-MM-dd")
        else:
            date = None
        cur.execute('''UPDATE todos
        SET name = ?, info = ?, color_id = ?, date = ?
        WHERE id = ?''',
                    (self.nameInput.text(), self.infoEdit.toPlainText(),
                     self.colorBox.currentIndex() + 1, date, self.id))
        self.con.commit()
        set_settings()

    def change_color(self):  # меняем цвет плана в его окне и в списке
        cur = self.con.cursor()
        color = cur.execute(
            "SELECT color_code FROM colors WHERE color_id = ?",
            (self.colorBox.currentIndex() + 1, )).fetchone()[0]
        self.spacer.setStyleSheet(f"background-color: rgb({color})")

    def go_back(self, event):  # закрываем окно
        self.close()

    def change_date(self):
        # в зависимости от статуса deadlineBox активируем или деактивируем поле ввода даты
        self.dateEdit.setEnabled(self.deadlineBox.isChecked())


class SettingsWindow(BaseWindow, Ui_SettingsWindow):  # окно настройки представленния списка
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.con = sqlite3.connect("notes_db.sqlite")
        self.closeButton.clicked.connect(self.close_settings)
        self.hideButton.clicked.connect(self.showMinimized)
        self.titleBar.mousePressEvent = self.titleBarMousePressEvent

        self.defaultButton.mousePressEvent = self.set_default
        self.doneButton.mousePressEvent = self.set_settings

    def close_settings(self):  # закрываем окно
        self.close()

    def set_default(self, event):
        # восстанавливаем исходные настройки представления (без фильтров)
        set_settings(refresh=True)

    def set_settings(self, event):  # применяем имеющиеся настройки представления
        color_id = self.colorBox.currentIndex()
        has_deadline = self.deadlineBox.isChecked()
        is_undone = self.undoneBox.isChecked()

        settings.change(color_id, has_deadline, is_undone)
        set_settings()


class CalendarWindow(BaseWindow, Ui_CalendarWindow):  # календарь для планов
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.con = sqlite3.connect("notes_db.sqlite")

        self.closeButton.clicked.connect(self.close_calendar)
        self.hideButton.clicked.connect(self.showMinimized)
        self.titleBar.mousePressEvent = self.titleBarMousePressEvent

        self.calendarWidget.currentPageChanged.connect(self.highlight_dates)
        self.calendarWidget.activated.connect(self.create_todo)
        self.calendarWidget.selectionChanged.connect(self.show_selected_todos)

        self.update_dates_list()
        self.highlight_format = QTextCharFormat()
        self.highlight_format.setBackground(Qt.lightGray)
        self.highlight_format.setForeground(self.palette().color(QPalette.HighlightedText))
        self.highlight_dates(
            self.calendarWidget.yearShown(), self.calendarWidget.monthShown())

        self.calendarWidget.setSelectedDate(dt.date.today())
        date = self.calendarWidget.selectedDate()
        # отображаем в списке только планы на выбранную дату (сегодня)
        settings.change(date=date)
        set_settings()

        self.show_selected_todos()

    def show_selected_todos(self):
        # отображаем в списке планов те, которые закреплены за выбранной в календаре датой
        date = self.calendarWidget.selectedDate()
        settings.change(date=date)
        set_settings()

    def update_dates_list(self):
        # актуализируем список дат
        cur = self.con.cursor()
        dates = cur.execute(
            "SELECT date FROM todos WHERE date IS NOT NULL").fetchall()
        self.dates = [date[0] for date in dates]  # список имеющихся дат

    def highlight_dates(self, year, month):
        #  выделяем в текущем листе календаря даты, к которым прикреплены планы

        def date_iter(year, month):
            #  вспомогательный метод, возвращающий итератор
            #  со всеми датами открытого на данный момент месяца
            for i in range(1, monthrange(year, month)[1] + 1):
                yield dt.date(year, month, i).strftime("%Y-%m-%d")

        for day in date_iter(year, month):
            if day in self.dates:
                self.calendarWidget.setDateTextFormat(
                    dt.datetime.strptime(day, "%Y-%m-%d"), self.highlight_format)

    def create_todo(self):
        #  создание нового плана при двойном нажатии на дату
        date = self.calendarWidget.selectedDate()
        todo = ToDoWindow(get_new_todo_id(), True, date)
        todo.show()
        self.dates.append(date.toString("yyyy-MM-dd"))
        self.highlight_dates(
            self.calendarWidget.yearShown(), self.calendarWidget.monthShown())

    def close_calendar(self):  # закрытие окна
        set_settings(refresh=True)
        self.close()


class ToDoSettings:
    #  служебный класс для хранения и изменения настроек представления
    def __init__(self):
        self.settings_are_active = False
        self.refresh()

    def refresh(self):  # возвращаем настройки представления по умолчанию
        self.params = [None, False, False, None]

    def change(self,
               color_index=None, has_deadline=None,
               is_undone=None, date=None):
        #  добавляем новые настройки представления к уже имеющимся
        parameters = [color_index, has_deadline, is_undone, date]
        for i in range(4):
            if parameters[i] is not None and parameters[i] != self.params[i]:
                self.params[i] = parameters[i]


class InfoWindow(BaseWindow, Ui_InfoWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.closeButton.clicked.connect(self.close_info)
        self.hideButton.clicked.connect(self.showMinimized)
        self.titleBar.mousePressEvent = self.titleBarMousePressEvent

        self.githubIcon.mousePressEvent = self.open_github

    def close_info(self):
        self.close()

    def open_github(self, event):
        webbrowser.open("https://github.com/danilasch/workraum")


if __name__ == '__main__':
    app = QApplication(sys.argv)
    notes = NotesMenu()
    todos = ToDoMenu()
    notes.show()
    settings = ToDoSettings()
    sys.exit(app.exec())
